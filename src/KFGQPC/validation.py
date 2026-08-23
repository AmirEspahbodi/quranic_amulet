"""Structural, Unicode, font, differential, and round-trip validation gates."""

from __future__ import annotations

from collections import Counter
from pathlib import Path
from zipfile import ZipFile
import subprocess
import unicodedata

from docx import Document
from lxml import etree

from . import SOURCE_SYMBOL_FONT, TARGET_FONT
from .ooxml import (
    NS,
    W,
    extract_segments,
    file_sha256,
    inspect_package,
    iter_paragraph_runs,
    load_font_resolver,
    package_member_hashes,
    parse_xml,
    run_text_nodes,
)
from .reference import reference_transform
from .rules import is_arabic_or_quranic, is_effective_delete, known_effective_source
from .transformer import transform_segment


STRUCTURAL_TAGS = (
    "p", "r", "t", "tbl", "tr", "tc", "sectPr", "bookmarkStart",
    "bookmarkEnd", "hyperlink", "drawing", "fldChar", "instrText", "txbxContent",
)


def _structure_counts(root: etree._Element) -> dict[str, int]:
    return {tag: len(root.findall(f".//w:{tag}", NS)) for tag in STRUCTURAL_TAGS}


def _direct_fonts(run: etree._Element) -> dict[str, str]:
    rfonts = run.find("./w:rPr/w:rFonts", NS)
    if rfonts is None:
        return {}
    return {etree.QName(key).localname: value for key, value in rfonts.attrib.items()}


def _orphan_marks(text: str) -> list[dict]:
    occurrences: list[dict] = []
    at_word_start = True
    for index, char in enumerate(text):
        if char.isspace():
            at_word_start = True
            continue
        if at_word_start and unicodedata.category(char).startswith("M"):
            occurrences.append({
                "index": index,
                "codepoint": f"U+{ord(char):04X}",
                "name": unicodedata.name(char, "UNNAMED"),
                "context": text[max(0, index - 8):index + 9].encode("unicode_escape").decode("ascii"),
            })
        at_word_start = False
    return occurrences


def _font_file(target_font: str) -> dict:
    try:
        result = subprocess.run(
            ["fc-match", "-f", "%{family}\n%{file}\n", target_font],
            check=True,
            capture_output=True,
            text=True,
            timeout=15,
        )
        lines = [line.strip() for line in result.stdout.splitlines() if line.strip()]
    except (OSError, subprocess.SubprocessError):
        return {"available": False, "family": None, "path": None, "cmap_checked": False}
    family = lines[0] if lines else None
    font_path = lines[1] if len(lines) > 1 else None
    exact = bool(family and target_font.lower() in {item.strip().lower() for item in family.split(",")})
    return {"available": exact, "family": family, "path": font_path if exact else None, "cmap_checked": False}


def _font_cmap_check(font_info: dict, codepoints: set[int]) -> dict:
    if not font_info["available"] or not font_info["path"]:
        return font_info
    try:
        from fontTools.ttLib import TTFont
        font = TTFont(font_info["path"], lazy=True)
        cmap: set[int] = set()
        for table in font["cmap"].tables:
            cmap.update(table.cmap)
        relevant = {cp for cp in codepoints if not chr(cp).isspace()}
        missing = sorted(relevant - cmap)
        font.close()
        return {
            **font_info,
            "cmap_checked": True,
            "required_codepoint_count": len(relevant),
            "missing_codepoints": [f"U+{cp:04X}" for cp in missing],
            "coverage_ok": not missing,
        }
    except Exception as exc:
        return {**font_info, "cmap_checked": False, "cmap_error": str(exc)}


def validate_output(
    input_path: Path,
    output_path: Path,
    transform_report: dict,
    *,
    original_input_sha256: str,
    target_font: str = TARGET_FONT,
    source_symbol_font: str = SOURCE_SYMBOL_FONT,
    deterministic_bytes_match: bool = True,
) -> dict:
    output_inspection = inspect_package(output_path)
    input_inspection = inspect_package(input_path)
    input_unchanged = file_sha256(input_path) == original_input_sha256
    try:
        document = Document(output_path)
        docx_library_reopen = True
        docx_paragraph_count = len(document.paragraphs)
    except Exception as exc:
        docx_library_reopen = False
        docx_paragraph_count = None
        docx_error = str(exc)

    input_member_hashes = package_member_hashes(input_path)
    output_member_hashes = package_member_hashes(output_path)
    modified_parts = set(transform_report["modified_parts"])
    unchanged_part_failures = sorted(
        name for name, digest in input_member_hashes.items()
        if name not in modified_parts and output_member_hashes.get(name) != digest
    )

    target_locations = {
        (item["part"], item["paragraph"], item["run"])
        for item in transform_report["target_run_locations"]
    }
    font_failures: list[dict] = []
    non_target_run_failures: list[dict] = []
    structural_failures: list[dict] = []
    target_texts: list[str] = []
    idempotence_failures: list[dict] = []
    oracle_failures: list[dict] = []
    output_codepoints: set[int] = set()

    with ZipFile(input_path) as source, ZipFile(output_path) as output:
        output_resolver = load_font_resolver(output)
        for part in output_inspection.story_parts:
            source_root = parse_xml(source.read(part), part)
            output_root = parse_xml(output.read(part), part)
            source_structure = _structure_counts(source_root)
            output_structure = _structure_counts(output_root)
            if source_structure != output_structure:
                structural_failures.append({
                    "part": part,
                    "input": source_structure,
                    "output": output_structure,
                })
            source_paragraphs = list(source_root.iter(f"{W}p"))
            output_paragraphs = list(output_root.iter(f"{W}p"))
            for paragraph_index, (source_p, output_p) in enumerate(zip(source_paragraphs, output_paragraphs)):
                source_runs = list(iter_paragraph_runs(source_p))
                output_runs = list(iter_paragraph_runs(output_p))
                for run_index, (source_run, output_run) in enumerate(zip(source_runs, output_runs)):
                    location = (part, paragraph_index, run_index)
                    output_text = "".join(node.text or "" for node in run_text_nodes(output_run))
                    if location in target_locations and output_text:
                        fonts = _direct_fonts(output_run)
                        required = {key: target_font for key in ("ascii", "hAnsi", "eastAsia", "cs")}
                        theme_keys = {key for key in fonts if key.endswith("Theme") or key == "cstheme"}
                        if any(fonts.get(key) != value for key, value in required.items()) or theme_keys:
                            font_failures.append({
                                "part": part,
                                "paragraph": paragraph_index,
                                "run": run_index,
                                "fonts": fonts,
                            })
                    elif location not in target_locations:
                        if etree.tostring(source_run) != etree.tostring(output_run):
                            non_target_run_failures.append({
                                "part": part,
                                "paragraph": paragraph_index,
                                "run": run_index,
                            })

            segments, _ = extract_segments(
                part, output_root, output_resolver,
                target_font=target_font,
                source_symbol_font=source_symbol_font,
            )
            for segment in segments:
                text = segment.text
                target_texts.append(text)
                output_codepoints.update(map(ord, text))
                rerun = transform_segment(segment, strict=True)
                if rerun.text != text:
                    idempotence_failures.append({
                        "part": part,
                        "paragraph": segment.paragraph_index,
                        "segment": segment.segment_index,
                    })
                oracle = reference_transform(text, surah_start=segment.surah_start, strict=True)
                if oracle != rerun.text:
                    oracle_failures.append({
                        "part": part,
                        "paragraph": segment.paragraph_index,
                        "segment": segment.segment_index,
                    })

    output_target_text = "".join(target_texts)
    persian_yeh_count = output_target_text.count("\u06cc")
    remaining_deletes = Counter(
        f"U+{ord(char):04X}" for char in output_target_text if is_effective_delete(char)
    )
    unknown_output = Counter(
        f"U+{ord(char):04X}" for char in output_target_text
        if is_arabic_or_quranic(char) and not known_effective_source(char)
    )
    orphan_marks = _orphan_marks(output_target_text)
    font_info = _font_cmap_check(_font_file(target_font), output_codepoints)

    checks = {
        "zip_crc_and_required_parts": output_inspection.crc_ok and output_inspection.required_parts_present,
        "docx_library_reopen": docx_library_reopen,
        "input_unchanged": input_unchanged,
        "non_modified_parts_byte_identical": not unchanged_part_failures,
        "modified_part_structure_preserved": not structural_failures,
        "non_target_runs_byte_identical": not non_target_run_failures,
        "all_surviving_target_runs_exact_font": not font_failures,
        "no_persian_yeh": persian_yeh_count == 0,
        "no_effective_delete_codepoint": not remaining_deletes,
        "no_unknown_arabic_quranic_codepoint": not unknown_output,
        "no_orphan_combining_mark": not orphan_marks,
        "no_unapproved_duplicate_alef": transform_report["unapproved_duplicate_alef"] == 0,
        "lineage_complete": bool(transform_report["lineage_complete"]),
        "independent_oracle_all_segments": all(
            item["independent_oracle_match"] for item in transform_report["segments"]
        ) and not oracle_failures,
        "idempotent_target_text": not idempotence_failures,
        "deterministic_package_bytes": deterministic_bytes_match,
        "no_unresolved_ambiguity": transform_report["unresolved_count"] == 0,
    }
    return {
        "input_sha256": input_inspection.sha256,
        "output_sha256": output_inspection.sha256,
        "output_package": output_inspection.__dict__,
        "docx_library_paragraph_count": docx_paragraph_count,
        "docx_library_error": None if docx_library_reopen else docx_error,
        "checks": checks,
        "all_required_checks_passed": all(checks.values()),
        "unchanged_part_failures": unchanged_part_failures,
        "structural_failures": structural_failures,
        "non_target_run_failures": non_target_run_failures,
        "font_failures": font_failures,
        "persian_yeh_count": persian_yeh_count,
        "remaining_delete_codepoints": dict(sorted(remaining_deletes.items())),
        "unknown_output_codepoints": dict(sorted(unknown_output.items())),
        "orphan_combining_marks": orphan_marks,
        "idempotence_failures": idempotence_failures,
        "oracle_failures": oracle_failures,
        "font_file": font_info,
        "target_output_character_count": len(output_target_text),
        "target_output_histogram": dict(sorted(Counter(f"U+{ord(c):04X}" for c in output_target_text).items())),
    }
