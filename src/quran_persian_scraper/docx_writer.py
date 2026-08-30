import logging
from pathlib import Path
from docx import Document
from src.quran_persian_scraper.models import Surah

logger = logging.getLogger(__name__)


class DocxWriter:
    """Saves surah translations to DOCX format adhering to strict formatting guidelines."""

    def __init__(self, output_dir: Path) -> None:
        self._output_dir = output_dir
        self._output_dir.mkdir(parents=True, exist_ok=True)

    def save_surah(self, surah: Surah) -> Path:
        """
        Creates a DOCX containing only the extracted Persian verses separated
        by exactly one space, with no extra characters, titles, or separators.
        """
        doc = Document()

        # Add single paragraph with the exact space-joined verses
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(surah.combined_translation)

        # Set RTL flow for Persian text display
        paragraph.paragraph_format.right_to_left = True

        file_path = self._output_dir / surah.formatted_filename
        doc.save(file_path)
        logger.info("Saved: %s (%d verses)", file_path.name, len(surah.verses))
        return file_path