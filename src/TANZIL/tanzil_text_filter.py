#!/usr/bin/env python3
"""
Quranic Text Filter: Tanzil Uthmani → Base/Initial Quran
Font: KFGQPC HAFS Uthmanic Script
Pipeline: Phase 1 (Contextual) → Phase 2 (Mapping) → Phase 3 (Deletion)

The program deliberately transforms complete paragraph strings rather than
individual Word runs.  Tanzil combining signs can be separated from their base
letters by run boundaries, so processing runs independently would be unsafe.
"""

from __future__ import annotations

import argparse
import re
import sys
import tempfile
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


# === CONFIGURATION ===
INPUT_DIR = "Quran_Surahs"
OUTPUT_DIR = "Quran_Surahs_herz"
FONT_NAME = "KFGQPC HAFS Uthmanic Script"
FONT_SIZE_PT = 18.0
ARABIC_LANGUAGE = "ar-SA"


# All documented Muqatta'at forms.  This is a generic orthographic safeguard,
# not Surah-specific processing.  It prevents the final maddah in forms such as
# alif-lam-meem from being mistaken for a consonantal maddah (Rule C.2).
_MUQATTAAT_FORMS = (
    "المص",  # alif-lam-meem-sad
    "كهيعص",  # kaf-ha-ya-ain-sad
    "المر",  # alif-lam-meem-ra
    "الم",  # alif-lam-meem
    "الر",  # alif-lam-ra
    "طسم",  # ta-seen-meem
    "طس",  # ta-seen
    "طه",  # ta-ha
    "يس",  # ya-seen
    "حم",  # ha-meem
    "عسق",  # ain-seen-qaf (the second part of ha-meem ain-seen-qaf)
    "ص",  # sad
    "ق",  # qaf
    "ن",  # noon
)

_ARABIC_BASE_RANGE = r"\u0621-\u064A"
_QURANIC_MARK_RANGE = r"\u064B-\u065F\u0670\u06D6-\u06ED"


def _marked_form_pattern(base_form: str) -> str:
    """Return a regex for a base form with optional Quranic marks per letter."""

    marks = f"[{_QURANIC_MARK_RANGE}]*"
    return "".join(re.escape(letter) + marks for letter in base_form)


_MUQATTAAT_RE = re.compile(
    rf"(?<![{_ARABIC_BASE_RANGE}{_QURANIC_MARK_RANGE}])"
    rf"(?:{'|'.join(_marked_form_pattern(x) for x in _MUQATTAAT_FORMS)})"
    rf"(?![{_ARABIC_BASE_RANGE}{_QURANIC_MARK_RANGE}])"
)


def _strip_muqattaat_maddah(match: re.Match[str]) -> str:
    """Strip both maddah codepoints only inside a recognized Muqatta'at form."""

    return re.sub(r"[\u0653\u06E4]", "", match.group(0))


def _collapse_double_alef(text: str) -> str:
    """Rule C.3: collapse every run of two or more simple Alefs."""

    return re.sub(r"\u0627{2,}", "\u0627", text)


def phase_1_contextual(text: str) -> str:
    """Apply Rules C.1 through C.8 using context-aware regular expressions."""

    # C.5 first: protect plural-verb waw + maddah + silent Alef as one unit.
    text = re.sub(r"\u0648\u0653\u0627", "\u0648\u0627", text)

    # C.1: the explicit alif-lam-maddah form plus every complete recognized
    # Muqatta'at token.  The token pass also strips a final meem/sad/qaf maddah.
    text = re.sub(r"\u0627\u0644\u0653", "\u0627\u0644", text)
    text = _MUQATTAAT_RE.sub(_strip_muqattaat_maddah, text)

    # C.1: strip maddah after a long vowel.  U+0670, U+06E5, and U+06E6
    # are included as Tanzil long-vowel proxies so later C.7/T.6/D.3 cannot
    # expose a stranded maddah after Phase 1.
    text = re.sub(
        r"(?<=[\u0648\u064A\u0627\u0670\u06E5\u06E6])\u0653",
        "",
        text,
    )

    # C.2: when an existing Alef follows, strip the maddah rather than creating
    # a second Alef.  Otherwise convert consonantal maddah to simple Alef.
    consonants = r"[\u0621-\u0647\u0649-\u064A\u0671]"
    text = re.sub(rf"(?<={consonants})\u0653(?=\u0627)", "", text)
    text = re.sub(rf"(?<={consonants})\u0653(?!\u0627)", "\u0627", text)

    # C.4: inverted damma before Alef contributes no additional Alef; in every
    # other specified context it becomes a simple Alef.
    text = re.sub(r"\u0657(?=\u0627)", "", text)
    text = re.sub(r"\u0657(?!\u0627)", "\u0627", text)

    # C.6: sequential dammatan glyph normalization.
    text = re.sub(r"\u065E", "\u064C", text)

    # C.7: dagger Alef.  The first two substitutions cover both collapse
    # conditions explicitly; the safe substitution handles all other cases.
    text = re.sub(r"(?<=\u0627)\u0670", "", text)
    text = re.sub(r"\u0670(?=\u0627)", "", text)
    text = re.sub(r"(?<!\u0627)\u0670(?!\u0627)", "\u0627", text)

    # C.8: small high madda follows the same long-vowel/consonant distinction.
    # By this point C.7 has converted a preceding dagger Alef to simple Alef.
    text = re.sub(
        r"(?<=[\u0648\u064A\u0627\u06E5\u06E6])\u06E4",
        "",
        text,
    )
    text = re.sub(rf"(?<={consonants})\u06E4(?=\u0627)", "", text)
    text = re.sub(rf"(?<={consonants})\u06E4(?!\u0627)", "\u0627", text)

    # Strip any remaining U+06E4 signs (such as Tanzil Sajdah markers on vocalized letters)
    text = re.sub(r"\u06E4", "", text)

    # C.3 is required after all Phase 1 Alef-producing rules.
    return _collapse_double_alef(text)

# === PHASE 2: Character Mapping ===
_PHASE_2_MAP = {
    "\u0671": "\u0627",  # T.1  Alef Wasla -> Alef
    "\u0622": "\u0627",  # T.2  Alef with Madda -> Alef
    "\u0625": "\u0627",  # T.3  Alef with Hamza Below -> Alef
    "\u0649": "\u0627",  # T.4  Alef Maksura -> Alef
    "\u06E3": "\u062B",  # T.5  Small Low Seen -> Tha
    "\u06E6": "\u064A",  # T.6  Small Yeh -> Arabic Yeh
    "\u06E8": "\u0646",  # T.7  Small High Noon -> Noon
    "\u06ED": "\u0645",  # T.8  Small Low Meem -> Meem
    "\u06E7": "\u064A",  # T.9  Small High Yeh -> Arabic Yeh
}
_PHASE_2_RE = re.compile("[" + re.escape("".join(_PHASE_2_MAP)) + "]")


def phase_2_mapping(text: str) -> str:
    """Apply Rules T.1 through T.9 and then re-run Rule C.3."""

    text = _PHASE_2_RE.sub(lambda match: _PHASE_2_MAP[match.group(0)], text)
    return _collapse_double_alef(text)


# === PHASE 3: Deletion ===
# Duplicate entries D.17, D.21, and D.22 intentionally collapse to their
# literal codepoints.  U+06EB and U+06EC are therefore listed once each.  The
# specification assigns D.22 to literal U+06EC; U+06E0 is not in its deletion
# list and is consequently preserved rather than silently inferred by name.
_PHASE_3_DELETION_CHARS = (
    "\u0651"  # D.1  Shadda
    "\u06E2"  # D.2  Small High Meem Isolated
    "\u06E5"  # D.3  Small Waw
    "\u06EA"  # D.4  Empty Centre Low Stop
    "\u06EB"  # D.5 / D.21
    "\u06EC"  # D.6 / D.22
    "\u06DD"  # D.7  End of Ayah
    "\u06DE"  # D.8  Rub el Hizb
    "\u06D4"  # D.9  Arabic Full Stop
    "\u06D6"  # D.10
    "\u06D7"  # D.11
    "\u06D8"  # D.12
    "\u06D9"  # D.13
    "\u06DA"  # D.14
    "\u06DB"  # D.15 / D.17
    "\u06DC"  # D.16
    "\u0640"  # D.19 Tatweel
)
_PHASE_3_DELETION_RE = re.compile(
    "[" + re.escape(_PHASE_3_DELETION_CHARS) + "]"
)

# D.18: Tanzil places an ayah-number run directly after the final Quranic word
# as well as between spaces.  A following Arabic/ASCII letter means the digits
# are embedded in a token and must be preserved.  ASCII-token adjacency is also
# protected.  This removes the attached verse-number suffixes in real Tanzil
# data without a blind global digit deletion.
_AYAH_NUMBER_RE = re.compile(
    r"(?<![0-9A-Za-z_\u0660-\u0669])"
    r"[\u0660-\u0669]+"
    r"(?![0-9A-Za-z_\u0660-\u0669\u0621-\u064A\u0671])"
)

# D.20 runs after Phase 2, so original qaf/sad-lam-Alef-Maksura signs have
# already become qaf/sad-lam-Alef.  Longest alternatives must be attempted first.
_PAUSE_MARK_ALTERNATION = r"(?:\u0642\u0644\u0627|\u0635\u0644\u0627|\u0644\u0627|\u0645|\u062C|\u0637|\u0642|\u0635)"
_PAUSE_MARK_RE = re.compile(
    rf"(?<=\s){_PAUSE_MARK_ALTERNATION}(?=\s)"
)
_PAUSE_MARK_AND_ONE_SEPARATOR_RE = re.compile(
    rf"(?<=\s){_PAUSE_MARK_ALTERNATION}(?=[ \t\u00A0]\S)[ \t\u00A0]"
)


def phase_3_deletion(text: str) -> str:
    """Apply Rules D.1 through D.22 without deleting letters inside words."""

    # Remove exactly one following separator with an isolated pause mark when
    # possible.  The preceding separator remains, so "word mark word" retains
    # the original single inter-word separation rather than becoming two spaces.
    text = _PAUSE_MARK_AND_ONE_SEPARATOR_RE.sub("", text)
    text = _PAUSE_MARK_RE.sub("", text)

    # Remove maximal ayah-number runs before deleting U+06DD; this also handles
    # inputs where the number immediately follows the end-of-ayah sign.
    text = _AYAH_NUMBER_RE.sub("", text)
    return _PHASE_3_DELETION_RE.sub("", text)


# === MAIN PIPELINE ===
def process_text(text: str) -> str:
    """Apply all three phases in the mandated order."""

    text = phase_1_contextual(text)
    text = phase_2_mapping(text)
    text = phase_3_deletion(text)
    # Final safety: collapse any remaining double Alefs.
    text = re.sub(r"\u0627{2,}", "\u0627", text)
    validate_output_text(text)
    return text


_FORBIDDEN_AFTER_PIPELINE = set(
    "\u0651\u0653\u0657\u065E\u0670\u06E4"
    + "".join(_PHASE_2_MAP)
    + _PHASE_3_DELETION_CHARS
)
_EXPLICIT_ALLOWED_CONTROLS = {"\u200C", "\u200D", "\u200E", "\u200F"}


def _format_codepoint(character: str) -> str:
    return f"U+{ord(character):04X}"


def validate_character_domain(text: str, *, label: str) -> None:
    """Reject characters that cannot belong to the requested output domain."""

    for offset, character in enumerate(text):
        codepoint = ord(character)
        if character == "\u06CC":
            raise ValueError(f"{label} contains forbidden Farsi Yeh U+06CC")
        if character.isspace():
            continue
        if character in _EXPLICIT_ALLOWED_CONTROLS:
            continue
        if 0x0600 <= codepoint <= 0x06FF:
            continue
        raise ValueError(
            f"{label} contains unsupported character "
            f"{_format_codepoint(character)} at character offset "
            f"{offset}"
        )


def validate_output_text(text: str) -> None:
    """Enforce the post-pipeline Unicode and collision invariants."""

    validate_character_domain(text, label="Transformed text")
    if "\u0627\u0627" in text:
        raise ValueError("Transformed text still contains a double Alef")

    remaining = sorted(set(text).intersection(_FORBIDDEN_AFTER_PIPELINE), key=ord)
    if remaining:
        codes = ", ".join(_format_codepoint(ch) for ch in remaining)
        raise ValueError(f"Pipeline left forbidden source/deletion characters: {codes}")


# === DOCX I/O ===
def _reject_unreadable_body_content(document: Document) -> None:
    """Fail rather than silently omit text python-docx cannot expose as paragraphs."""

    if document.tables:
        raise ValueError(
            "Input contains table text; paragraph preservation would be ambiguous"
        )

    body = document.element.body
    unsupported_nodes = body.xpath(
        ".//w:txbxContent//w:t | .//w:sdt//w:t | .//w:ins//w:t | .//w:del//w:delText"
    )
    if any((node.text or "") for node in unsupported_nodes):
        raise ValueError(
            "Input contains text boxes, content controls, or tracked revisions; "
            "refusing to omit or reinterpret sacred text"
        )


def read_docx(path: str | Path) -> list[str]:
    """Read every main-document paragraph, preserving empty paragraphs and order."""

    source_path = Path(path)
    if source_path.suffix.lower() != ".docx":
        raise ValueError(f"Input must be a .docx file: {source_path}")
    if not source_path.is_file():
        raise FileNotFoundError(f"Input file does not exist: {source_path}")

    document = Document(str(source_path))
    _reject_unreadable_body_content(document)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    if not paragraphs or not any(text.strip() for text in paragraphs):
        raise ValueError("Input document contains no Quranic paragraph text")

    for index, text in enumerate(paragraphs, start=1):
        validate_character_domain(text, label=f"Input paragraph {index}")
    return paragraphs


def _get_or_add_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _set_font_xml(r_pr, font_size_pt: float) -> None:
    """Set all Word font slots, including the complex-script slot."""

    r_fonts = _get_or_add_child(r_pr, "w:rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), FONT_NAME)
    r_fonts.set(qn("w:hint"), "cs")

    size_half_points = str(int(round(font_size_pt * 2)))
    sz_cs = _get_or_add_child(r_pr, "w:szCs")
    sz_cs.set(qn("w:val"), size_half_points)


def _format_paragraph(paragraph, font_size_pt: float) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    p_pr = paragraph._p.get_or_add_pPr()
    _get_or_add_child(p_pr, "w:bidi")

    # An explicit paragraph-mark font also covers intentionally empty paragraphs.
    p_r_pr = _get_or_add_child(p_pr, "w:rPr")
    _set_font_xml(p_r_pr, font_size_pt)


def _format_run(run, font_size_pt: float) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(font_size_pt)
    run.font.complex_script = True
    run.font.rtl = True
    run.font.no_proof = True

    r_pr = run._element.get_or_add_rPr()
    _set_font_xml(r_pr, font_size_pt)
    language = _get_or_add_child(r_pr, "w:lang")
    language.set(qn("w:val"), ARABIC_LANGUAGE)
    language.set(qn("w:bidi"), ARABIC_LANGUAGE)


def write_docx(
    path: str | Path,
    paragraphs: Sequence[str],
    *,
    font_size_pt: float = FONT_SIZE_PT,
) -> None:
    """Write a new, RTL DOCX with one explicitly formatted run per paragraph."""

    if not 6.0 <= font_size_pt <= 96.0:
        raise ValueError("Font size must be between 6 and 96 points")

    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(font_size_pt)
    normal.font.complex_script = True
    normal.font.rtl = True
    normal.paragraph_format.space_after = Pt(0)
    normal_r_pr = normal.element.get_or_add_rPr()
    _set_font_xml(normal_r_pr, font_size_pt)

    for text in paragraphs:
        paragraph = document.add_paragraph(style=normal)
        _format_paragraph(paragraph, font_size_pt)
        run = paragraph.add_run(text)
        _format_run(run, font_size_pt)

    document.save(str(output_path))


def verify_written_docx(
    path: str | Path,
    expected_paragraphs: Sequence[str],
) -> None:
    """Reopen the package and verify exact text, structure, RTL, and font slots."""

    document = Document(str(path))
    actual_paragraphs = [paragraph.text for paragraph in document.paragraphs]
    if actual_paragraphs != list(expected_paragraphs):
        mismatch = next(
            (
                index
                for index, (actual, expected) in enumerate(
                    zip(actual_paragraphs, expected_paragraphs), start=1
                )
                if actual != expected
            ),
            min(len(actual_paragraphs), len(expected_paragraphs)) + 1,
        )
        raise ValueError(
            "DOCX round-trip text mismatch at paragraph "
            f"{mismatch}; expected {len(expected_paragraphs)} paragraphs, "
            f"found {len(actual_paragraphs)}"
        )

    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        p_pr = paragraph._p.pPr
        if p_pr is None or p_pr.find(qn("w:bidi")) is None:
            raise ValueError(f"Paragraph {paragraph_index} is missing RTL metadata")
        if not paragraph.runs:
            raise ValueError(f"Paragraph {paragraph_index} has no formatted run")

        for run_index, run in enumerate(paragraph.runs, start=1):
            r_pr = run._element.rPr
            r_fonts = r_pr.rFonts if r_pr is not None else None
            if r_fonts is None:
                raise ValueError(
                    f"Paragraph {paragraph_index}, run {run_index} has no font record"
                )
            for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
                actual_font = r_fonts.get(qn(f"w:{attribute}"))
                if actual_font != FONT_NAME:
                    raise ValueError(
                        f"Paragraph {paragraph_index}, run {run_index}, "
                        f"font slot {attribute} is {actual_font!r}, not {FONT_NAME!r}"
                    )


def transform_docx(
    input_path: str | Path,
    output_path: str | Path,
    *,
    font_size_pt: float = FONT_SIZE_PT,
    overwrite: bool = False,
) -> tuple[int, int, int]:
    """Transform one DOCX atomically and return paragraph/input/output counts."""

    source = Path(input_path)
    destination = Path(output_path)
    if source.resolve() == destination.resolve():
        raise ValueError("Output must be a new file; it cannot overwrite the input")
    if destination.exists() and not overwrite:
        raise FileExistsError(
            f"Output already exists: {destination} (use --force to replace it)"
        )

    original_paragraphs = read_docx(source)
    transformed_paragraphs = [process_text(text) for text in original_paragraphs]

    destination.parent.mkdir(parents=True, exist_ok=True)
    temporary_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(
            prefix=f".{destination.stem}.",
            suffix=".docx",
            dir=destination.parent,
            delete=False,
        ) as temporary_file:
            temporary_path = Path(temporary_file.name)

        write_docx(
            temporary_path,
            transformed_paragraphs,
            font_size_pt=font_size_pt,
        )
        verify_written_docx(temporary_path, transformed_paragraphs)
        temporary_path.replace(destination)
        temporary_path = None
    finally:
        if temporary_path is not None and temporary_path.exists():
            temporary_path.unlink()

    input_characters = sum(len(text) for text in original_paragraphs)
    output_characters = sum(len(text) for text in transformed_paragraphs)
    return len(original_paragraphs), input_characters, output_characters


# === BUILT-IN REGRESSION CHECKS ===
def _expect(actual: str, expected: str, rule: str) -> None:
    if actual != expected:
        raise AssertionError(f"{rule} failed: expected {expected!r}, got {actual!r}")


def run_self_tests() -> None:
    """Exercise every transformation family without requiring a DOCX file."""

    phase_1_cases = (
        ("\u0641\u064a\u0653", "\u0641\u064a", "C.1 long-vowel Yeh"),
        ("\u0633\u064f\u0648\u0653\u0621\u064e", "\u0633\u064f\u0648\u0621\u064e", "C.1 long-vowel Waw"),
        ("\u0627\u0644\u0653\u0645\u0653", "\u0627\u0644\u0645", "C.1 Muqatta'at"),
        ("\u0628\u0653", "\u0628\u0627", "C.2 consonantal maddah"),
        ("\u0628\u0653\u0627", "\u0628\u0627", "C.2 existing Alef"),
        ("\u0627\u0627\u0627", "\u0627", "C.3 double Alef"),
        ("\u0645\u064e\u0631\u064e\u0636\u0657\u0627", "\u0645\u064e\u0631\u064e\u0636\u0627", "C.4 inverted damma"),
        ("\u0642\u064e\u0627\u0644\u064f\u0648\u0653\u0627\u0652", "\u0642\u064e\u0627\u0644\u064f\u0648\u0627\u0652", "C.5 plural Waw"),
        ("\u0631\u064e\u062d\u0650\u064a\u0645\u065e", "\u0631\u064e\u062d\u0650\u064a\u0645\u064c", "C.6 sequential dammatan"),
        ("\u0647\u064e\u0670\u0630\u064e\u0627", "\u0647\u064e\u0627\u0630\u064e\u0627", "C.7 dagger Alef"),
        ("\u0627\u0670", "\u0627", "C.7 preceding Alef"),
        ("\u0628\u0670\u0627", "\u0628\u0627", "C.7 following Alef"),
        ("\u064a\u06e4", "\u064a", "C.8 long-vowel madda"),
        ("\u0628\u06e4", "\u0628\u0627", "C.8 consonantal madda"),
        ("\u064a\u064e\u0633\u0652\u062c\u064f\u062f\u064f\u0648\u0646\u064e\u06e4", "\u064a\u064e\u0633\u0652\u062c\u064f\u062f\u064f\u0648\u0646\u064e", "C.8 sajdah mark with fatha"),
        ("\u0648\u064e\u0671\u0633\u0652\u062c\u064f\u062f\u0652\u06e4", "\u0648\u064e\u0671\u0633\u0652\u062c\u064f\u062f\u0652", "C.8 sajdah mark with sukun"),
    )
    for source, expected, rule in phase_1_cases:
        _expect(phase_1_contextual(source), expected, rule)

    # Every known Muqatta'at base form must survive intact when each letter is
    # marked with maddah; only the combining maddah signs may disappear.
    for base_form in _MUQATTAAT_FORMS:
        marked_form = "".join(letter + "\u0653" for letter in base_form)
        _expect(
            phase_1_contextual(marked_form),
            base_form,
            f"C.1 Muqatta'at form {base_form!r}",
        )

    phase_2_cases = (
        ("\u0671", "\u0627", "T.1"),
        ("\u0622", "\u0627", "T.2"),
        ("\u0625", "\u0627", "T.3"),
        ("\u0649", "\u0627", "T.4"),
        ("\u06E3", "\u062B", "T.5"),
        ("\u06E6", "\u064A", "T.6"),
        ("\u06E8", "\u0646", "T.7"),
        ("\u06ED", "\u0645", "T.8"),
        ("\u06E7", "\u064A", "T.9"),
        ("\u0671\u0627", "\u0627", "C.3 after Phase 2"),
    )
    for source, expected, rule in phase_2_cases:
        _expect(phase_2_mapping(source), expected, rule)

    _expect(process_text("\u0627\u0644\u0653\u0645\u0653\u0661"), "\u0627\u0644\u0645", "full pipeline")
    _expect(process_text("\u0649\u0670\u0653"), "\u0627", "dagger-Alef/maddah/Phase-2 collision")
    _expect(process_text("\u06E6\u0653"), "\u064A", "small-Yeh/maddah/Phase-2 collision")
    _expect(process_text("\u0643\u0644\u0645\u0629\u0661\u0662 \u0643\u0644\u0645\u0629"), "\u0643\u0644\u0645\u0629 \u0643\u0644\u0645\u0629", "D.18 ayah number")
    _expect(process_text("\u0643\u0661\u0644\u0645\u0629"), "\u0643\u0661\u0644\u0645\u0629", "D.18 embedded digit")
    _expect(process_text("\u0643\u0644\u0645\u0629 \u0645 \u0643\u0644\u0645\u0629"), "\u0643\u0644\u0645\u0629 \u0643\u0644\u0645\u0629", "D.20 pause mark")
    _expect(process_text("\u0644\u064e\u0627"), "\u0644\u064e\u0627", "D.20 vocalized word")
    _expect(process_text("\u0645\u0651\u0640\u06D6"), "\u0645", "D.1/D.10/D.19")
    _expect(process_text(_PHASE_3_DELETION_CHARS), "", "all literal Phase 3 deletions")

    kept = "\u064e\u064f\u0650\u064b\u064c\u064d\u0652\u06e1\u0654\u0655\u0656\u06e9\u0623\u200c\u200d\u200e\u200f"
    _expect(process_text(kept), kept, "required pass-through characters")

def _build_argument_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description=(
            f"Process all .docx Surah files from '{INPUT_DIR}' and save "
            f"filtered output to '{OUTPUT_DIR}' using font '{FONT_NAME}'."
        )
    )
    parser.add_argument(
        "--input-dir",
        type=Path,
        default=Path(INPUT_DIR),
        help=f"input directory containing .docx files (default: {INPUT_DIR})",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path(OUTPUT_DIR),
        help=f"output directory for filtered .docx files (default: {OUTPUT_DIR})",
    )
    parser.add_argument(
        "--font-size",
        type=float,
        default=FONT_SIZE_PT,
        metavar="PT",
        help=f"font size in points (default: {FONT_SIZE_PT:g})",
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="replace existing files in the output directory",
    )
    parser.add_argument(
        "--self-test",
        action="store_true",
        help="run built-in Unicode regression checks and exit",
    )
    return parser


def main(argv: Sequence[str] | None = None) -> int:
    parser = _build_argument_parser()
    args = parser.parse_args(argv)

    if args.self_test:
        run_self_tests()
        print("All Unicode pipeline self-tests passed.")
        return 0

    input_dir = args.input_dir
    output_dir = args.output_dir

    if not input_dir.exists() or not input_dir.is_dir():
        raise FileNotFoundError(f"Input directory does not exist: {input_dir.resolve()}")

    docx_files = sorted(
        [f for f in input_dir.glob("*.docx") if not f.name.startswith("~$")]
    )

    if not docx_files:
        print(f"No .docx files found in {input_dir.resolve()}")
        return 0

    output_dir.mkdir(parents=True, exist_ok=True)
    print(f"Found {len(docx_files)} files to process in '{input_dir}'.")

    success_count = 0
    for file_path in docx_files:
        target_path = output_dir / file_path.name
        try:
            p_count, in_chars, out_chars = transform_docx(
                file_path,
                target_path,
                font_size_pt=args.font_size,
                overwrite=args.force,
            )
            print(
                f"[OK] {file_path.name} -> {target_path.name} "
                f"({p_count} paragraphs | {in_chars} -> {out_chars} chars)"
            )
            success_count += 1
        except Exception as err:
            print(f"[ERROR] Failed to process {file_path.name}: {err}", file=sys.stderr)

    print(f"\nFinished: Successfully processed {success_count}/{len(docx_files)} files.")
    return 0 if success_count == len(docx_files) else 1


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except (AssertionError, FileNotFoundError, OSError, ValueError) as error:
        print(f"Error: {error}", file=sys.stderr)
        raise SystemExit(1)
