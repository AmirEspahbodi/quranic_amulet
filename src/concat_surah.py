# -*- coding: utf-8 -*-
"""
ادغام فایل‌های Word قرآنی در یک فایل Word واحد با اعمال اجباری قلم عثمان طه.

- فایل‌ها بر اساس عدد ابتدای نام فایل مرتب می‌شوند.
- متن هر فایل استخراج و با یک جداکننده ادغام می‌شود.
- شماره منازل ۷ گانه قرآن با کاراکترهای عربی پیش از سوره‌های مربوطه درج می‌شود.
- فونت 'KFGQPC HAFS Uthmanic Script' در تمام لایه‌های OpenXML (DocDefaults، Styles و Runs)
  برای سازگاری کامل با Microsoft Word 2024 اعمال می‌شود.
"""

import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run


# ----------------------------- تنظیمات -----------------------------

SOURCE_DIR = Path(r"surahs") # آدرس پوشه را طبق لاگ شما تصحیح کردم
OUTPUT_FILE = Path(r"holly_quran.docx")
EXTENSIONS = {".docx", ".doc"}
PERIOD = "."
EXPECTED_COUNT = 114

# نام فونت اختصاصی مصحف
TARGET_FONT = "KFGQPC HAFS Uthmanic Script"
# کد زبان عربی برای جلوگیری از خطایابی نامناسب در Word 2024
LOCALE_ARABIC = "ar-SA"

# دیکشنری نگاشت نقطه شروع منازل قرآن به اعداد عربی جهت پشتیبانی در فونت عثمان طه
# اعداد یونیکد عربی در فونت KFGQPC دقیقاً با استایل قرآن رندر می‌شوند
MANZIL_STARTS = {
    1: "١",   # منزل اول: الفاتحه تا النساء
    5: "٢",   # منزل دوم: المائده تا التوبه
    10: "٣",  # منزل سوم: یونس تا النحل
    17: "٤",  # منزل چهارم: الاسراء تا الفرقان
    26: "٥",  # منزل پنجم: الشعراء تا یس
    37: "٦",  # منزل ششم: الصافات تا الحجرات
    50: "٧"   # منزل هفتم: ق تا الناس
}

# ------------------------------------------------------------------


def get_file_number(file_path: Path) -> Optional[int]:
    """استخراج شماره ترتیبی سوره از نام فایل."""
    match = re.match(r"^\s*(\d+)", file_path.name)
    if not match:
        match = re.search(r"(\d+)", file_path.name)

    return int(match.group(1)) if match else None


def normalize_text(text: str) -> str:
    """
    پاک‌سازی نویسه‌های کنترلی غیراستاندارد با حفظ علائم اعراب، تنوین،
    نمادهای وقوف و علائم انتهای آیه در خط عثمان طه.
    """
    if not text:
        return ""

    # حذف کاراکتر شکست سلول جدول در خروجی COM
    text = text.replace("\x07", " ")

    # حذف نویسه‌های کنترلی ASCII بدون دست‌کاری کدهای یونیکد عربی/قرآنی
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)

    # تبدیل فضاهای خالی تکراری به یک فاصله بدون حذف نیم‌فاصله یا علائم قرآنی
    text = re.sub(r"[^\S\r\n]+", " ", text)
    text = re.sub(r"\s+", " ", text)

    return text.strip()


def iter_block_items(doc: Document):
    """پیمایش ترتیبی تمام عناصر محتوایی بدنه سند (پاراگراف‌ها و جداول)."""
    def iter_children(elm):
        for child in elm.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, doc)
            elif child.tag == qn("w:tbl"):
                yield Table(child, doc)
            elif child.tag.endswith("sdt"):
                # پشتیبانی از Structured Document Tags (Content Controls)
                yield from iter_children(child)

    yield from iter_children(doc.element.body)


def extract_docx_text(file_path: Path) -> str:
    """استخراج تمیز متن از فایل‌های .docx شامل پاراگراف‌ها و سلول‌های جدول."""
    try:
        doc = Document(str(file_path))
    except Exception as exc:
        print(f"خطا در خواندن فایل {file_path.name}: {exc}")
        return ""

    parts: List[str] = []

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = normalize_text(block.text)
            if text:
                parts.append(text)
        elif isinstance(block, Table):
            for row in block.rows:
                row_texts = [normalize_text(c.text) for c in row.cells if normalize_text(c.text)]
                if row_texts:
                    parts.append(" ".join(row_texts))

    return normalize_text(" ".join(parts))


def extract_doc_texts_with_word_com(file_paths: List[Path]) -> Dict[Path, str]:
    """استخراج محتوای فایل‌های باینری قدیمی (.doc) از طریق Word Automation."""
    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:
        raise RuntimeError(
            "برای خواندن فایل‌های .doc قدیمی نیاز به pywin32 است:\n"
            "pip install pywin32"
        ) from exc

    result: Dict[Path, str] = {}
    pythoncom.CoInitialize()
    word = None

    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False

        for file_path in file_paths:
            try:
                doc = word.Documents.Open(
                    FileName=str(file_path.resolve()),
                    ConfirmConversions=False,
                    ReadOnly=True,
                    AddToRecentFiles=False,
                )
                text = doc.Content.Text
                result[file_path] = normalize_text(text)
                doc.Close(SaveChanges=False)
            except Exception as e:
                print(f"خطا در خواندن فایل {file_path.name} با Word COM: {e}")
                result[file_path] = ""
    finally:
        if word:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()

    return result


def apply_force_font_to_run(run: Run, font_name: str):
    """
    اعمال اجباری قلم قرآنی در سطح Run.
    تنظیم صریح ویژگی‌های اسکریپت پیچیده (w:cs) و حذف تم‌های پیش‌فرض Word 2024.
    """
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()

    # تنظیم یا ایجاد تگ w:rFonts
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)

    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)

    # پاک‌سازی پیوندهای Theme که باعث جایگزینی فونت در Word 2024 می‌شوند
    theme_attrs = [
        "asciiTheme",
        "hAnsiTheme",
        "cstheme",
        "eastAsiaTheme",
    ]
    for attr in theme_attrs:
        key = qn(f"w:{attr}")
        if key in rFonts.attrib:
            del rFonts.attrib[key]

    # مشخص کردن جهت راست‌به‌چپ در سطح نویسه
    rtl = rPr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        rPr.append(rtl)
    rtl.set(qn("w:val"), "1")

    # تنظیم زبان به عربی جهت اتصال صحیح حروفی در موتور متنی ورد
    lang = rPr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rPr.append(lang)
    lang.set(qn("w:val"), LOCALE_ARABIC)
    lang.set(qn("w:bidi"), LOCALE_ARABIC)


def apply_force_font_to_document_defaults(doc: Document, font_name: str):
    """
    تغییر تنظیمات پیش‌فرض کل سند (docDefaults) و استایل Normal
    به فونت هدف برای جلوگیری از هرگونه فال‌بک نرم‌افزاری.
    """
    # 1. به‌روزرسانی استایل پیش‌فرض 'Normal'
    normal_style = doc.styles["Normal"]
    normal_style.font.name = font_name
    normal_rPr = normal_style.element.get_or_add_rPr()

    rFonts_normal = normal_rPr.find(qn("w:rFonts"))
    if rFonts_normal is None:
        rFonts_normal = OxmlElement("w:rFonts")
        normal_rPr.insert(0, rFonts_normal)

    rFonts_normal.set(qn("w:ascii"), font_name)
    rFonts_normal.set(qn("w:hAnsi"), font_name)
    rFonts_normal.set(qn("w:cs"), font_name)
    rFonts_normal.set(qn("w:eastAsia"), font_name)

    for attr in ["asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"]:
        key = qn(f"w:{attr}")
        if key in rFonts_normal.attrib:
            del rFonts_normal.attrib[key]

    # 2. به‌روزرسانی w:docDefaults در بخش styles.xml
    styles_element = doc.styles.element
    doc_defaults = styles_element.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_element.insert(0, doc_defaults)

    r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
    if r_pr_default is None:
        r_pr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(r_pr_default)

    r_pr = r_pr_default.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        r_pr_default.append(r_pr)

    r_fonts_default = r_pr.find(qn("w:rFonts"))
    if r_fonts_default is None:
        r_fonts_default = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts_default)

    r_fonts_default.set(qn("w:ascii"), font_name)
    r_fonts_default.set(qn("w:hAnsi"), font_name)
    r_fonts_default.set(qn("w:cs"), font_name)
    r_fonts_default.set(qn("w:eastAsia"), font_name)

    for attr in ["asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"]:
        key = qn(f"w:{attr}")
        if key in r_fonts_default.attrib:
            del r_fonts_default.attrib[key]


def set_paragraph_rtl(paragraph: Paragraph):
    """تنظیم مشخصات راست‌به‌چپ (Bidi) و چیدمان پاراگراف."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    pPr = paragraph._p.get_or_add_pPr()

    # تعریف جهت Bidi
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    bidi.set(qn("w:val"), "1")

    # تراز راست OpenXML
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "right")


def main():
    if not SOURCE_DIR.is_dir():
        raise FileNotFoundError(f"پوشه ورودی یافت نشد: {SOURCE_DIR.resolve()}")

    files: List[Tuple[Optional[int], Path]] = []
    for path in SOURCE_DIR.iterdir():
        if path.is_file() and path.suffix.lower() in EXTENSIONS:
            files.append((get_file_number(path), path))

    if not files:
        raise FileNotFoundError(
            f"هیچ فایلی با پسوندهای {sorted(EXTENSIONS)} در پوشه پیدا نشد: {SOURCE_DIR.resolve()}"
        )

    # بررسی فایل‌های فاقد شماره
    no_number_files = [path for num, path in files if num is None]
    if no_number_files:
        print("هشدار: این فایل‌ها عدد ابتدایی ندارند و در انتهای پردازش قرار می‌گیرند:")
        for path in no_number_files:
            print(f"  - {path.name}")

    # اعتبارسنجی تکرار و فقدان شماره‌ها
    numbers = [num for num, _ in files if num is not None]
    seen = set()
    duplicates = set()
    for num in numbers:
        if num in seen:
            duplicates.add(num)
        seen.add(num)

    if duplicates:
        print("هشدار: شماره‌های تکراری مشاهده شد:", sorted(duplicates))

    missing = [n for n in range(1, EXPECTED_COUNT + 1) if n not in seen]
    if missing:
        print("هشدار: این شماره سوره‌ها در پوشه غایب هستند:", missing)

    # مرتب‌سازی منطقی
    files.sort(
        key=lambda item: (
            item[0] is None,
            item[0] if item[0] is not None else float("inf"),
            item[1].name.lower(),
        )
    )

    # خواندن تجمیعی فایل‌های .doc قدیمی
    doc_paths = [p for _, p in files if p.suffix.lower() == ".doc"]
    doc_texts = extract_doc_texts_with_word_com(doc_paths) if doc_paths else {}

    merged_parts: List[str] = []

    for num, path in files:
        label = num if num is not None else "?"
        print(f"در حال پردازش سوره {label}: {path.name}")

        if path.suffix.lower() == ".doc":
            text = doc_texts.get(path, "")
        else:
            text = extract_docx_text(path)

        if not text:
            print(f"هشدار: فایل {path.name} خالی یا نامعتبر بود.")
            continue

        # اگر سوره، سوره ابتدای یک منزل است، ابتدا شماره عربی منزل را به لیست اضافه می‌کنیم
        if num in MANZIL_STARTS:
            merged_parts.append(MANZIL_STARTS[num])

        # سپس متن سوره را اضافه می‌کنیم
        merged_parts.append(text)

    if not merged_parts:
        raise ValueError("هیچ متنی برای درج در فایل نهایی تولید نشد.")

    # با این جداکننده، چه بین "شماره منزل و سوره" و چه بین "سوره و سوره"
    # دقیقا الگو [space + dot + space] رعایت می‌شود.
    separator = f" {PERIOD} "
    
    # متصل کردن تمامی عناصر داخل لیست با جداکننده
    final_text = separator.join(merged_parts)
    
    # مطابق خواسته شما در انتها: "سوره اخر باید یک فاصله و بعد نقطه داشته باشه"
    final_text += f" {PERIOD}"

    # ایجاد سند نهایی با اعمال سخت‌گیرانه تایپوگرافی
    final_doc = Document()
    apply_force_font_to_document_defaults(final_doc, TARGET_FONT)

    paragraph = final_doc.add_paragraph()
    set_paragraph_rtl(paragraph)

    # افزودن متن در قالب Run و الزام قطعی فونت عثمان طه
    run = paragraph.add_run(final_text)
    apply_force_font_to_run(run, TARGET_FONT)

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    final_doc.save(str(OUTPUT_FILE))

    print("\nفایل نهایی با قلم KFGQPC HAFS Uthmanic Script ساخته شد:")
    print(OUTPUT_FILE.resolve())


if __name__ == "__main__":
    main()
