from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from quran_docx_filter import SOURCE_SYMBOL_FONT, TARGET_FONT
from quran_docx_filter.ooxml import RunMetadata, SourceToken, TextNodeRef, TextSegment


def _metadata(run_index: int, *, raised: bool = False, ambiguous: bool = False, target: bool = True) -> RunMetadata:
    if ambiguous:
        direct = (("ascii", SOURCE_SYMBOL_FONT), ("cs", "Ambiguous Arabic"), ("hAnsi", SOURCE_SYMBOL_FONT))
        resolved = direct
    else:
        direct = (("ascii", SOURCE_SYMBOL_FONT), ("cs", TARGET_FONT), ("hAnsi", SOURCE_SYMBOL_FONT))
        resolved = direct
    return RunMetadata(
        part="word/document.xml",
        paragraph_index=0,
        run_index=run_index,
        direct_fonts=direct,
        resolved_fonts=resolved,
        run_style=None,
        paragraph_style=None,
        vert_align="superscript" if raised else None,
        position=None,
        rtl=True,
        paragraph_bidi=True,
        target=target,
    )


@pytest.fixture
def make_segment():
    def factory(
        text: str,
        *,
        split_every_char: bool = False,
        split_at: set[int] | None = None,
        raised_indices: set[int] | None = None,
        ambiguous_indices: set[int] | None = None,
        surah_start: bool = False,
    ) -> TextSegment:
        raised_indices = raised_indices or set()
        ambiguous_indices = ambiguous_indices or set()
        split_at = split_at or set()
        groups: list[tuple[int, int]] = []
        start = 0
        for index in range(1, len(text)):
            if split_every_char or index in split_at:
                groups.append((start, index))
                start = index
        groups.append((start, len(text)))
        tokens: list[SourceToken] = []
        nodes: list[TextNodeRef] = []
        for run_index, (start, end) in enumerate(groups):
            # Split again when formatting evidence changes inside a requested group.
            sub_start = start
            signatures = [
                (i in raised_indices, i in ambiguous_indices) for i in range(start, end)
            ]
            boundaries = [
                start + offset for offset in range(1, len(signatures))
                if signatures[offset] != signatures[offset - 1]
            ] + [end]
            for boundary in boundaries:
                actual_run = len(nodes)
                raised = sub_start in raised_indices
                ambiguous = sub_start in ambiguous_indices
                run = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
                node = etree.SubElement(run, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                node.text = text[sub_start:boundary]
                metadata = _metadata(actual_run, raised=raised, ambiguous=ambiguous)
                ref = TextNodeRef(node, run, metadata, 0, node.text or "")
                nodes.append(ref)
                for char_in_node, source_index in enumerate(range(sub_start, boundary)):
                    ref.source_indices.append(source_index)
                    tokens.append(SourceToken(text[source_index], source_index, ref, char_in_node))
                sub_start = boundary
        return TextSegment("word/document.xml", 0, 0, tokens, nodes, surah_start)
    return factory


def set_run_font(run, font: str) -> None:
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), font)


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    path = tmp_path / "sample.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph._p.get_or_add_pPr().append(etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi"))
    for text in ("ب", "\u0651", "\u06dd", "ت"):
        run = paragraph.add_run(text)
        set_run_font(run, TARGET_FONT)
    title = paragraph.add_run(" metadata: 41-54")
    set_run_font(title, "Tahoma")
    table = document.add_table(rows=1, cols=1)
    table_run = table.cell(0, 0).paragraphs[0].add_run("ف\u06cc")
    set_run_font(table_run, TARGET_FONT)
    document.save(path)
    return path


@pytest.fixture
def corpus_docx_paths() -> list[Path]:
    root = Path(__file__).resolve().parents[1]
    return sorted(
        path for path in root.glob("*.docx")
        if not path.name.endswith(".filtered.docx")
    )
