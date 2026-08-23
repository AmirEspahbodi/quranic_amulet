from __future__ import annotations

from hashlib import sha256
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree

from quran_docx_filter import SOURCE_SYMBOL_FONT, TARGET_FONT
from quran_docx_filter.ooxml import inspect_package
from quran_docx_filter.transformer import transform_docx_to_bytes
from quran_docx_filter.validation import validate_output


def test_docx_roundtrip_reopens_and_preserves_input(sample_docx: Path, tmp_path: Path):
    before = sha256(sample_docx.read_bytes()).hexdigest()
    result = transform_docx_to_bytes(sample_docx, strict=True)
    output = tmp_path / "roundtrip.filtered.docx"
    output.write_bytes(result.package_bytes)
    Document(output)
    assert sha256(sample_docx.read_bytes()).hexdigest() == before
    assert inspect_package(output).crc_ok


def test_docx_output_is_deterministic(sample_docx: Path):
    first = transform_docx_to_bytes(sample_docx, strict=True)
    second = transform_docx_to_bytes(sample_docx, strict=True)
    assert first.package_bytes == second.package_bytes


def test_table_text_is_transformed_and_non_target_prose_unchanged(sample_docx: Path, tmp_path: Path):
    result = transform_docx_to_bytes(sample_docx, strict=True)
    output = tmp_path / "table.filtered.docx"
    output.write_bytes(result.package_bytes)
    document = Document(output)
    assert document.tables[0].cell(0, 0).text == "في"
    assert "metadata: 41-54" in document.paragraphs[0].text


def test_full_validation_gates_pass_on_fixture(sample_docx: Path, tmp_path: Path):
    input_sha = sha256(sample_docx.read_bytes()).hexdigest()
    first = transform_docx_to_bytes(sample_docx, strict=True)
    second = transform_docx_to_bytes(sample_docx, strict=True)
    output = tmp_path / "validated.filtered.docx"
    output.write_bytes(first.package_bytes)
    validation = validate_output(
        sample_docx,
        output,
        first.report,
        original_input_sha256=input_sha,
        deterministic_bytes_match=first.package_bytes == second.package_bytes,
    )
    assert validation["all_required_checks_passed"], validation


def test_target_runs_have_only_required_font(sample_docx: Path, tmp_path: Path):
    result = transform_docx_to_bytes(sample_docx, strict=True)
    output = tmp_path / "font.filtered.docx"
    output.write_bytes(result.package_bytes)
    with ZipFile(output) as package:
        root = etree.fromstring(package.read("word/document.xml"))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    target_texts = {"ب", "ت", "ف", "ي", " "}
    for run in root.xpath(".//w:r[w:t]", namespaces=ns):
        text = "".join(run.xpath(".//w:t/text()", namespaces=ns))
        if "metadata" in text:
            continue
        rfonts = run.find("./w:rPr/w:rFonts", namespaces=ns)
        if not text or rfonts is None:
            continue
        values = {
            etree.QName(key).localname: value for key, value in rfonts.attrib.items()
            if etree.QName(key).localname in {"ascii", "hAnsi", "eastAsia", "cs"}
        }
        assert values == {key: TARGET_FONT for key in ("ascii", "hAnsi", "eastAsia", "cs")}


def test_target_text_inside_textbox_is_processed(sample_docx: Path, tmp_path: Path):
    # Inject a minimal txbxContent subtree; the transformer must process the
    # inner paragraph once and never flatten it into the outer paragraph.
    injected = tmp_path / "textbox.docx"
    with ZipFile(sample_docx) as source, ZipFile(injected, "w") as target:
        for info in source.infolist():
            data = source.read(info.filename)
            if info.filename == "word/document.xml":
                root = etree.fromstring(data)
                ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                body = root.find(f"{{{ns}}}body")
                outer = etree.Element(f"{{{ns}}}p")
                pict_run = etree.SubElement(outer, f"{{{ns}}}r")
                txbx = etree.SubElement(pict_run, f"{{{ns}}}txbxContent")
                inner = etree.SubElement(txbx, f"{{{ns}}}p")
                run = etree.SubElement(inner, f"{{{ns}}}r")
                rpr = etree.SubElement(run, f"{{{ns}}}rPr")
                rfonts = etree.SubElement(rpr, f"{{{ns}}}rFonts")
                for key in ("ascii", "hAnsi", "eastAsia", "cs"):
                    rfonts.set(f"{{{ns}}}{key}", TARGET_FONT)
                text = etree.SubElement(run, f"{{{ns}}}t")
                text.text = "ب\u0651ت"
                body.insert(len(body) - 1, outer)
                data = etree.tostring(root, encoding="UTF-8", xml_declaration=True, standalone=True)
            target.writestr(info, data)
    result = transform_docx_to_bytes(injected, strict=True)
    output = tmp_path / "textbox.filtered.docx"
    output.write_bytes(result.package_bytes)
    with ZipFile(output) as package:
        text = etree.fromstring(package.read("word/document.xml")).xpath(
            "string(.//*[local-name()='txbxContent']//*[local-name()='t'])"
        )
    assert text == "بت"

